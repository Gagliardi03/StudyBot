from typing import Dict, Any, Optional, List, Tuple
import io
import os
import tempfile
from fastapi import UploadFile


# Importações condicionais (instaladas apenas se necessário)
try:
    import PyPDF2

    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    import docx

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


async def extract_text_from_pdf(file_content: bytes) -> Tuple[str, Dict[str, Any]]:
    """
    Extrai texto e metadados de um arquivo PDF

    Args:
        file_content: Conteúdo binário do arquivo PDF

    Returns:
        Texto extraído e dicionário de metadados
    """
    if not HAS_PYPDF2:
        raise ImportError(
            "PyPDF2 não está instalado. Execute 'pip install PyPDF2' para habilitar extração de PDF."
        )

    metadata = {}
    text_content = ""

    try:
        # Lê o PDF a partir dos bytes
        pdf_file = io.BytesIO(file_content)
        reader = PyPDF2.PdfReader(pdf_file)

        # Extrai metadados
        if reader.metadata:
            for key, value in reader.metadata.items():
                if key.startswith("/"):
                    key = key[1:]  # Remove a barra inicial
                metadata[key] = value

        # Extrai número de páginas
        metadata["pageCount"] = len(reader.pages)

        # Extrai texto de cada página
        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_content += f"--- Página {page_num + 1} ---\n{page_text}\n\n"

        return text_content, metadata
    except Exception as e:
        return f"Erro ao processar PDF: {str(e)}", {"error": str(e)}


async def extract_text_from_docx(file_content: bytes) -> Tuple[str, Dict[str, Any]]:
    """
    Extrai texto e metadados de um arquivo DOCX

    Args:
        file_content: Conteúdo binário do arquivo DOCX

    Returns:
        Texto extraído e dicionário de metadados
    """
    if not HAS_DOCX:
        raise ImportError(
            "python-docx não está instalado. Execute 'pip install python-docx' para habilitar extração de DOCX."
        )

    metadata = {}
    text_content = ""

    try:
        # Salva o conteúdo em um arquivo temporário (python-docx requer um arquivo)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
            temp_path = temp_file.name
            temp_file.write(file_content)

        # Abre o documento
        doc = docx.Document(temp_path)

        # Extrai metadados do documento
        core_properties = doc.core_properties
        metadata = {
            "author": core_properties.author,
            "created": core_properties.created,
            "modified": core_properties.modified,
            "title": core_properties.title,
            "subject": core_properties.subject,
            "paragraphCount": len(doc.paragraphs),
        }

        # Extrai texto de cada parágrafo
        for para in doc.paragraphs:
            if para.text:
                text_content += para.text + "\n"

        # Remove o arquivo temporário
        os.unlink(temp_path)

        return text_content, metadata
    except Exception as e:
        # Garante que o arquivo temporário seja removido em caso de erro
        if "temp_path" in locals() and os.path.exists(temp_path):
            os.unlink(temp_path)
        return f"Erro ao processar DOCX: {str(e)}", {"error": str(e)}


async def extract_text_from_upload(
    upload_file: UploadFile,
) -> Tuple[str, Dict[str, Any]]:
    """
    Extrai texto de um arquivo enviado com base em seu tipo

    Args:
        upload_file: Arquivo enviado pelo usuário

    Returns:
        Texto extraído e dicionário de metadados
    """
    # Lê o conteúdo do arquivo
    content = await upload_file.read()

    # Determina o tipo de arquivo pela extensão
    filename = upload_file.filename.lower()

    if filename.endswith(".pdf"):
        return await extract_text_from_pdf(content)
    elif filename.endswith((".docx", ".doc")):
        return await extract_text_from_docx(content)
    elif filename.endswith((".txt", ".md", ".rtf")):
        # Para arquivos de texto, simplesmente decodifica o conteúdo
        try:
            text_content = content.decode("utf-8")
        except UnicodeDecodeError:
            # Tenta outras codificações se UTF-8 falhar
            try:
                text_content = content.decode("latin-1")
            except:
                text_content = "[Erro ao decodificar o conteúdo do arquivo]"

        return text_content, {
            "contentType": upload_file.content_type,
            "size": len(content),
        }
    else:
        return f"Tipo de arquivo não suportado: {filename}", {
            "error": "Tipo de arquivo não suportado"
        }


async def get_document_chunks(
    text: str, chunk_size: int = 1000, overlap: int = 200
) -> List[str]:
    """
    Divide um documento em chunks menores para processamento pelo LLM

    Args:
        text: Texto completo do documento
        chunk_size: Tamanho máximo de cada chunk em caracteres
        overlap: Quantidade de sobreposição entre chunks

    Returns:
        Lista de chunks de texto
    """
    chunks = []

    # Se o texto for menor que o tamanho do chunk, retorne-o diretamente
    if len(text) <= chunk_size:
        return [text]

    # Divide o texto em parágrafos
    paragraphs = text.split("\n")
    current_chunk = ""

    for para in paragraphs:
        # Se o parágrafo for maior que o tamanho do chunk, divida-o
        if len(para) > chunk_size:
            # Adiciona o chunk atual se não estiver vazio
            if current_chunk:
                chunks.append(current_chunk)
                current_chunk = ""

            # Divide o parágrafo em frases
            sentences = para.split(". ")
            temp_chunk = ""

            for sentence in sentences:
                if len(temp_chunk) + len(sentence) + 2 <= chunk_size:
                    temp_chunk += sentence + ". "
                else:
                    if temp_chunk:
                        chunks.append(temp_chunk)
                    temp_chunk = sentence + ". "

            if temp_chunk:
                current_chunk = temp_chunk
        else:
            # Se adicionar o parágrafo exceder o tamanho do chunk, finalize o chunk atual
            if len(current_chunk) + len(para) + 1 > chunk_size:
                chunks.append(current_chunk)
                current_chunk = para + "\n"
            else:
                current_chunk += para + "\n"

    # Adiciona o último chunk se não estiver vazio
    if current_chunk:
        chunks.append(current_chunk)

    # Adiciona sobreposição entre os chunks
    if overlap > 0 and len(chunks) > 1:
        chunks_with_overlap = [chunks[0]]

        for i in range(1, len(chunks)):
            prev_chunk = chunks[i - 1]
            curr_chunk = chunks[i]

            # Adiciona o final do chunk anterior ao início do atual
            if len(prev_chunk) > overlap:
                overlap_text = prev_chunk[-overlap:]
                chunks_with_overlap.append(overlap_text + curr_chunk)
            else:
                chunks_with_overlap.append(curr_chunk)

        return chunks_with_overlap

    return chunks
